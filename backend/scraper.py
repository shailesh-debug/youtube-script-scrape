from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from statistics import median
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from googleapiclient.discovery import build
from groq import Groq

from .config import Settings
from .models import ContentType


MAX_COMMENTS_PER_VIDEO = 100

THEME_MAP = {
    "Humour / Reactions": ["lol", "haha", "funny", "bro", "bruh"],
    "Scepticism / Doubt": ["not true", "wrong", "false", "doubt", "source", "proof"],
    "Dice / Game Talk": ["dice", "roll", "game", "luck", "odds", "probability"],
    "Hype / Amazement": ["insane", "crazy", "wild", "wow", "mind blown", "no way"],
    "Requests": ["can you", "please do", "next video", "more", "part 2"],
    "Personal Stories": ["i was", "i have", "happened to me", "my friend"],
    "Agreement": ["facts", "true", "exactly", "so true", "accurate"],
}

BLUE = RGBColor(0x1A, 0x73, 0xE8)
LIGHT_BLUE_HEX = "dbe9f8"
DARK_GREY = RGBColor(0x40, 0x40, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def parse_channel_input(channel_url: str) -> str:
    value = channel_url.strip()
    if not value:
        raise ValueError("Channel URL is required.")

    if value.startswith("@"):
        return value[1:].strip("/")

    if "youtube.com" not in value and "/" not in value:
        return value.lstrip("@")

    parsed = urlparse(value if "://" in value else f"https://{value}")
    parts = [part for part in parsed.path.split("/") if part]
    for part in parts:
        if part.startswith("@"):
            return part[1:]

    if len(parts) >= 2 and parts[0] in {"channel", "c", "user"}:
        return parts[1]

    raise ValueError("Could not parse the channel handle or id from that URL.")


def parse_iso8601_duration(duration: str) -> int:
    match = re.match(r"PT(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?", duration)
    if not match:
        return 0
    hours = int(match.group(1) or 0)
    minutes = int(match.group(2) or 0)
    seconds = int(match.group(3) or 0)
    return hours * 3600 + minutes * 60 + seconds


def clean_transcript(text: str) -> str:
    text = re.sub(r"\[.*?\]", "", text)
    return re.sub(r"\s+", " ", text).strip()


def get_channel_id(youtube, handle_or_id: str) -> str:
    if handle_or_id.startswith("UC"):
        return handle_or_id

    resp = youtube.search().list(
        part="snippet",
        q=handle_or_id,
        type="channel",
        maxResults=1,
    ).execute()
    items = resp.get("items", [])
    if not items:
        raise ValueError(f"Channel not found: {handle_or_id}")
    return items[0]["snippet"]["channelId"]


def get_video_ids_via_ytdlp(handle_or_id: str, content_type: ContentType, scan_size: int) -> list[str]:
    paths = []
    if content_type in {ContentType.shorts, ContentType.both}:
        paths.append("shorts")
    if content_type in {ContentType.videos, ContentType.both}:
        paths.append("videos")

    ids: list[str] = []
    channel_ref = f"channel/{handle_or_id}" if handle_or_id.startswith("UC") else f"@{handle_or_id}"
    for path in paths:
        url = f"https://www.youtube.com/{channel_ref}/{path}"
        result = subprocess.run(
            [
                sys.executable,
                "-m",
                "yt_dlp",
                "--flat-playlist",
                "--playlist-end",
                str(scan_size),
                "--print",
                "id",
                "--no-warnings",
                url,
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        ids.extend(line.strip() for line in result.stdout.splitlines() if line.strip())

    return list(dict.fromkeys(ids))[:scan_size]


def get_video_ids_via_api(youtube, channel_id: str, content_type: ContentType, scan_size: int) -> list[str]:
    ch = youtube.channels().list(part="contentDetails", id=channel_id).execute()
    playlist_id = ch["items"][0]["contentDetails"]["relatedPlaylists"]["uploads"]

    video_ids: list[str] = []
    next_page = None
    while len(video_ids) < scan_size:
        pl = youtube.playlistItems().list(
            part="contentDetails",
            playlistId=playlist_id,
            maxResults=min(50, scan_size - len(video_ids)),
            pageToken=next_page,
        ).execute()
        video_ids.extend(item["contentDetails"]["videoId"] for item in pl.get("items", []))
        next_page = pl.get("nextPageToken")
        if not next_page:
            break

    if content_type == ContentType.videos:
        return video_ids[:scan_size]

    filtered: list[str] = []
    for i in range(0, len(video_ids), 50):
        batch = video_ids[i : i + 50]
        details = youtube.videos().list(part="contentDetails", id=",".join(batch)).execute()
        for item in details.get("items", []):
            duration = parse_iso8601_duration(item["contentDetails"]["duration"])
            if content_type == ContentType.shorts and duration <= 180:
                filtered.append(item["id"])
            elif content_type == ContentType.both:
                filtered.append(item["id"])
    return filtered[:scan_size]


def fetch_video_metadata(youtube, video_ids: list[str]) -> list[dict]:
    results = []
    for i in range(0, len(video_ids), 50):
        batch = video_ids[i : i + 50]
        resp = youtube.videos().list(
            part="snippet,statistics,contentDetails",
            id=",".join(batch),
        ).execute()
        for item in resp.get("items", []):
            stats = item.get("statistics", {})
            results.append(
                {
                    "id": item["id"],
                    "title": item["snippet"]["title"],
                    "views": int(stats.get("viewCount", 0)),
                    "likes": int(stats.get("likeCount", 0)),
                    "comment_count": int(stats.get("commentCount", 0)),
                    "published": item["snippet"]["publishedAt"][:10],
                    "url": f"https://www.youtube.com/watch?v={item['id']}",
                    "transcript": "",
                    "comments": [],
                    "comment_summary": {},
                    "vs_baseline": 0.0,
                }
            )
    return results


def select_viral_videos(videos: list[dict], count: int) -> tuple[list[dict], dict]:
    if not videos:
        return [], {"median_views": 0, "threshold": 0, "note": "No videos found."}

    baseline = int(median(video["views"] for video in videos))
    threshold = int(baseline * 1.5)
    selected = [video for video in videos if video["views"] >= threshold]
    note = "Strong viral outliers found."

    if len(selected) < count:
        threshold = int(baseline * 1.1)
        selected = [video for video in videos if video["views"] >= threshold]
        note = "Relaxed viral threshold to include enough results."

    if len(selected) < count:
        selected = videos[:]
        note = "No strong viral outliers found; showing best performers."

    selected.sort(key=lambda video: video["views"], reverse=True)
    for video in selected:
        video["vs_baseline"] = video["views"] / baseline if baseline else 0

    return selected[:count], {
        "median_views": baseline,
        "threshold": threshold,
        "note": note,
    }


def _write_cookies_file(tmpdir: str, cookies: str) -> Path | None:
    if not cookies.strip():
        return None

    cookie_path = Path(tmpdir) / "youtube_cookies.txt"
    cookie_path.write_text(cookies.strip() + "\n", encoding="utf-8")
    return cookie_path


def _yt_dlp_failure_message(stderr: str) -> str:
    text = re.sub(r"\s+", " ", stderr).strip()
    if "Sign in to confirm" in text or "not a bot" in text:
        return "YouTube blocked anonymous audio download. Add YOUTUBE_COOKIES on Render."
    if not text:
        return "audio download failed"
    return text[:240]


def download_audio(video_id: str, tmpdir: str, cookies: str = "") -> tuple[Path | None, str]:
    output = Path(tmpdir) / "audio.%(ext)s"
    cookie_path = _write_cookies_file(tmpdir, cookies)
    cmd = [
        sys.executable,
        "-m",
        "yt_dlp",
        "--extract-audio",
        "--audio-format",
        "mp3",
        "--audio-quality",
        "5",
        "--no-playlist",
        "--no-warnings",
        "-o",
        str(output),
        f"https://www.youtube.com/watch?v={video_id}",
    ]
    if cookie_path:
        cmd[3:3] = ["--cookies", str(cookie_path)]

    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=180,
    )
    files = [Path(tmpdir) / name for name in os.listdir(tmpdir)]
    ignored_names = {"youtube_cookies.txt"}
    audio_path = next((path for path in files if path.is_file() and path.name not in ignored_names), None)
    return audio_path, _yt_dlp_failure_message(result.stderr)


def transcribe_with_groq(video_id: str, groq_api_key: str, youtube_cookies: str = "") -> str:
    if not groq_api_key:
        return "[No transcript - Groq API key is not configured]"

    tmpdir = tempfile.mkdtemp()
    try:
        audio_path, failure_message = download_audio(video_id, tmpdir, youtube_cookies)
        if not audio_path:
            return f"[No transcript - {failure_message}]"

        client = Groq(api_key=groq_api_key)
        with audio_path.open("rb") as audio:
            transcription = client.audio.transcriptions.create(
                file=(audio_path.name, audio.read()),
                model="whisper-large-v3",
                response_format="text",
                language="en",
            )
        return clean_transcript(str(transcription))
    except Exception as exc:
        return f"[No transcript - {exc}]"
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def get_top_comments(youtube, video_id: str, max_comments: int = MAX_COMMENTS_PER_VIDEO) -> list[dict]:
    comments = []
    try:
        next_page = None
        while len(comments) < max_comments:
            resp = youtube.commentThreads().list(
                part="snippet",
                videoId=video_id,
                order="relevance",
                maxResults=min(100, max_comments - len(comments)),
                pageToken=next_page,
            ).execute()
            for item in resp.get("items", []):
                top = item["snippet"]["topLevelComment"]["snippet"]
                comments.append({"text": top["textDisplay"], "likes": int(top.get("likeCount", 0))})
            next_page = resp.get("nextPageToken")
            if not next_page:
                break
    except Exception:
        pass
    return comments


def summarise_comments(comments: list[dict]) -> dict:
    theme_counts = {theme: 0 for theme in THEME_MAP}
    for comment in comments:
        lower = comment["text"].lower()
        for theme, keywords in THEME_MAP.items():
            if any(keyword in lower for keyword in keywords):
                theme_counts[theme] += 1

    return {
        "themes": {theme: count for theme, count in theme_counts.items() if count > 0},
        "top5": sorted(comments, key=lambda comment: comment["likes"], reverse=True)[:5],
    }


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)


def _heading(doc: Document, text: str, level: int):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = BLUE
    return heading


def build_docx(videos: list[dict], output_path: Path, channel_name: str, scan_summary: dict) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(channel_name)
    title_run.font.size = Pt(34)
    title_run.font.bold = True
    title_run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Viral YouTube Content Analysis")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = DARK_GREY

    generated = datetime.now().strftime("%d %b %Y")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Top {len(videos)} selected | Median: {scan_summary['median_views']:,} views | "
        f"Threshold: {scan_summary['threshold']:,} views | Generated: {generated}"
    )
    meta_run.font.size = Pt(11)
    meta_run.font.italic = True
    meta_run.font.color.rgb = DARK_GREY
    doc.add_page_break()

    _heading(doc, "Overview", level=1)
    note = doc.add_paragraph(scan_summary["note"])
    note.runs[0].font.italic = True

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["#", "Title", "Views", "vs Median", "Published"]):
        cell.text = text
        _set_cell_bg(cell, "1a73e8")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)

    for idx, video in enumerate(videos, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = video["title"]
        row[2].text = f"{video['views']:,}"
        row[3].text = f"{video['vs_baseline']:.1f}x"
        row[4].text = video["published"]
        if idx % 2 == 0:
            for cell in row:
                _set_cell_bg(cell, LIGHT_BLUE_HEX)
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    for idx, video in enumerate(videos, 1):
        _heading(doc, f"#{idx} {video['title']}", level=1)
        stats = doc.add_paragraph()
        stats_run = stats.add_run(
            f"{video['views']:,} views | {video['likes']:,} likes | "
            f"{video['comment_count']:,} comments | {video['published']} | {video['url']}"
        )
        stats_run.font.size = Pt(9)
        stats_run.font.italic = True
        stats_run.font.color.rgb = DARK_GREY

        _add_horizontal_rule(doc)
        _heading(doc, "Script / Transcript", level=2)
        transcript = doc.add_paragraph(video["transcript"] or "[No transcript]")
        for run in transcript.runs:
            run.font.size = Pt(10)

        _add_horizontal_rule(doc)
        _heading(doc, "Comment Summary", level=2)
        summary = video.get("comment_summary", {})
        themes = summary.get("themes", {})
        top5 = summary.get("top5", [])

        if themes:
            doc.add_paragraph("Comment themes detected:", style="Intense Quote")
            theme_paragraph = doc.add_paragraph()
            for theme, count in sorted(themes.items(), key=lambda item: -item[1]):
                run = theme_paragraph.add_run(f"{theme}: {count} comments\n")
                run.font.size = Pt(10)
        else:
            doc.add_paragraph("No theme matches found.")

        if top5:
            doc.add_paragraph("Top liked comments:", style="Intense Quote")
            for comment in top5:
                paragraph = doc.add_paragraph(style="List Bullet")
                run = paragraph.add_run(f"\"{comment['text'][:300]}\" - {comment['likes']:,} likes")
                run.font.size = Pt(9)
                run.font.italic = True
        else:
            paragraph = doc.add_paragraph("Comments not available for this video.")
            paragraph.runs[0].font.size = Pt(9)
            paragraph.runs[0].font.italic = True

        if idx < len(videos):
            doc.add_page_break()

    doc.save(output_path)


def run_scrape_job(settings: Settings, channel_url: str, count: int, content_type: ContentType, progress) -> dict:
    if not settings.youtube_api_key:
        raise ValueError("YouTube API key is not configured.")

    channel_ref = parse_channel_input(channel_url)
    scan_size = min(max(count * 6, 20), settings.max_scan_size)
    youtube = build("youtube", "v3", developerKey=settings.youtube_api_key)

    progress("scanning", 5, f"Scanning up to {scan_size} videos from {channel_ref}.")
    video_ids = get_video_ids_via_ytdlp(channel_ref, content_type, scan_size)
    if not video_ids:
        channel_id = get_channel_id(youtube, channel_ref)
        video_ids = get_video_ids_via_api(youtube, channel_id, content_type, scan_size)
    if not video_ids:
        raise ValueError("No videos found for that channel.")

    progress("filtering", 20, f"Fetched {len(video_ids)} video ids. Loading metadata.")
    videos = fetch_video_metadata(youtube, video_ids)
    videos.sort(key=lambda video: video["views"], reverse=True)
    selected, scan_summary = select_viral_videos(videos, count)

    for index, video in enumerate(selected, 1):
        percent = 25 + int((index - 1) / max(len(selected), 1) * 55)
        progress("transcribing", percent, f"Transcribing video {index} of {len(selected)}.")
        video["transcript"] = transcribe_with_groq(video["id"], settings.groq_api_key, settings.youtube_cookies)
        comments = get_top_comments(youtube, video["id"])
        video["comments"] = comments
        video["comment_summary"] = summarise_comments(comments)

    progress("building_report", 85, "Building DOCX report.")
    safe_channel = re.sub(r"[^A-Za-z0-9_-]+", "_", channel_ref).strip("_") or "channel"
    output_path = settings.output_dir / f"{safe_channel}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    build_docx(selected, output_path, channel_ref, scan_summary)

    return {
        "file_path": str(output_path),
        "summary": {
            "channel": channel_ref,
            "videos_scanned": len(videos),
            "videos_selected": len(selected),
            "median_views": scan_summary["median_views"],
            "viral_threshold": scan_summary["threshold"],
            "note": scan_summary["note"],
        },
    }
