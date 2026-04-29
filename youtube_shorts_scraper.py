"""
YouTube Shorts Scraper
Scrapes all Shorts from a YouTube channel and produces a formatted .docx report.
"""

import os
import re
import time
import subprocess
import sys
from datetime import datetime

from googleapiclient.discovery import build
import tempfile
from faster_whisper import WhisperModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants
from backend.config import get_settings

# ─── CONFIG ──────────────────────────────────────────────────────────────────

API_KEY      = get_settings().youtube_api_key
CHANNEL_HANDLE = "LoadedDiceShorts"
OUTPUT_FILE  = "LoadedDiceShorts_Shorts_Report.docx"
MAX_COMMENTS_PER_VIDEO = 100

THEME_MAP = {
    "😂 Humour / Reactions":  ["lol", "haha", "😂", "😭", "funny", "bro", "bruh", "💀"],
    "🤔 Scepticism / Doubt":  ["not true", "wrong", "false", "doubt", "source", "proof"],
    "🎲 Dice / Game Talk":    ["dice", "roll", "game", "luck", "odds", "probability"],
    "🙌 Hype / Amazement":    ["insane", "crazy", "wild", "wow", "mind blown", "no way"],
    "🙏 Requests":            ["can you", "please do", "next video", "more", "part 2"],
    "💬 Personal Stories":    ["i was", "i have", "happened to me", "my friend"],
    "🔥 Agreement":           ["facts", "true", "exactly", "so true", "accurate"],
}

# ─── HELPERS ─────────────────────────────────────────────────────────────────

def parse_iso8601_duration(duration: str) -> int:
    """Convert ISO 8601 duration (PT1M30S) to seconds."""
    match = re.match(
        r'PT(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?', duration
    )
    if not match:
        return 0
    hours   = int(match.group(1) or 0)
    minutes = int(match.group(2) or 0)
    seconds = int(match.group(3) or 0)
    return hours * 3600 + minutes * 60 + seconds


def clean_transcript(text: str) -> str:
    """Remove [Music], [Applause] and similar noise."""
    text = re.sub(r'\[.*?\]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


# ─── STEP 1: GET SHORTS IDS ──────────────────────────────────────────────────

def get_shorts_ids_via_ytdlp(handle: str) -> list[str]:
    """Primary method: scrape /shorts tab with yt-dlp."""
    url = f"https://www.youtube.com/@{handle}/shorts"
    print(f"[yt-dlp] Scraping {url} …")
    try:
        result = subprocess.run(
            [
                sys.executable, "-m", "yt_dlp",
                "--flat-playlist", "--print", "id",
                "--no-warnings",
                url,
            ],
            capture_output=True, text=True, timeout=120
        )
        ids = [line.strip() for line in result.stdout.splitlines() if line.strip()]
        print(f"[yt-dlp] Found {len(ids)} video IDs")
        return ids
    except Exception as e:
        print(f"[yt-dlp] Failed: {e}")
        return []


def get_channel_id(youtube, handle: str) -> str:
    """Resolve @handle → channel ID via YouTube API."""
    resp = youtube.search().list(
        part="snippet", q=handle, type="channel", maxResults=1
    ).execute()
    items = resp.get("items", [])
    if not items:
        raise ValueError(f"Channel not found: {handle}")
    return items[0]["snippet"]["channelId"]


def get_shorts_ids_via_api(youtube, channel_id: str) -> list[str]:
    """Fallback: fetch uploads playlist and filter ≤180 s."""
    print("[API] Fetching uploads playlist …")
    ch = youtube.channels().list(
        part="contentDetails", id=channel_id
    ).execute()
    playlist_id = ch["items"][0]["contentDetails"]["relatedPlaylists"]["uploads"]

    video_ids = []
    next_page = None
    while True:
        pl = youtube.playlistItems().list(
            part="contentDetails",
            playlistId=playlist_id,
            maxResults=50,
            pageToken=next_page,
        ).execute()
        batch = [item["contentDetails"]["videoId"] for item in pl.get("items", [])]
        video_ids.extend(batch)
        next_page = pl.get("nextPageToken")
        if not next_page:
            break

    print(f"[API] Total uploads: {len(video_ids)}. Filtering Shorts (≤180 s) …")

    # Batch into 50s and filter by duration
    shorts_ids = []
    for i in range(0, len(video_ids), 50):
        batch = video_ids[i:i+50]
        details = youtube.videos().list(
            part="contentDetails",
            id=",".join(batch),
        ).execute()
        for item in details.get("items", []):
            dur = parse_iso8601_duration(item["contentDetails"]["duration"])
            if dur <= 180:
                shorts_ids.append(item["id"])

    print(f"[API] Found {len(shorts_ids)} Shorts")
    return shorts_ids


def get_all_shorts(youtube, handle: str) -> list[str]:
    """Orchestrate: try yt-dlp first, fall back to API."""
    ids = get_shorts_ids_via_ytdlp(handle)
    if ids:
        return ids
    print("[Fallback] Using YouTube Data API …")
    channel_id = get_channel_id(youtube, handle)
    return get_shorts_ids_via_api(youtube, channel_id)


# ─── STEP 2: FETCH VIDEO METADATA ────────────────────────────────────────────

def fetch_video_metadata(youtube, video_ids: list[str]) -> list[dict]:
    """Batch-fetch title, views, likes, comments, date for all video IDs."""
    results = []
    for i in range(0, len(video_ids), 50):
        batch = video_ids[i:i+50]
        resp = youtube.videos().list(
            part="snippet,statistics,contentDetails",
            id=",".join(batch),
        ).execute()
        for item in resp.get("items", []):
            stats = item.get("statistics", {})
            results.append({
                "id":            item["id"],
                "title":         item["snippet"]["title"],
                "views":         int(stats.get("viewCount", 0)),
                "likes":         int(stats.get("likeCount", 0)),
                "comment_count": int(stats.get("commentCount", 0)),
                "published":     item["snippet"]["publishedAt"][:10],
                "url":           f"https://www.youtube.com/shorts/{item['id']}",
                "transcript":    "",
                "comments":      [],
                "comment_summary": {},
            })
        print(f"[Metadata] Fetched {min(i+50, len(video_ids))} / {len(video_ids)}")
    return results


# ─── STEP 3: FETCH TRANSCRIPTS (via yt-dlp audio + faster-whisper) ───────────

print("Loading Whisper model (tiny) …")
_whisper_model = WhisperModel("base", device="cuda", compute_type="float16")
print("Whisper model ready ✅")


def get_transcript(video_id: str) -> str:
    """Download audio via yt-dlp and transcribe with faster-whisper."""
    url = f"https://www.youtube.com/shorts/{video_id}"
    tmpdir = tempfile.mkdtemp()
    try:
        audio_out = os.path.join(tmpdir, "audio")
        cmd = [
            sys.executable, "-m", "yt_dlp",
            "--extract-audio", "--audio-format", "mp3", "--audio-quality", "5",
            "--no-playlist", "--no-warnings",
            "--cookies-from-browser", "firefox",
            "-o", audio_out + ".%(ext)s",
            url,
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

        audio_files = [
            os.path.join(tmpdir, f) for f in os.listdir(tmpdir)
            if os.path.isfile(os.path.join(tmpdir, f))
        ]
        if not audio_files:
            return "[No transcript — audio download failed]"

        segments, _ = _whisper_model.transcribe(audio_files[0], language="en")
        text = " ".join(seg.text.strip() for seg in segments)
        return clean_transcript(text)

    except Exception as e:
        return f"[No transcript — {e}]"
    finally:
        # Clean up temp audio files
        import shutil
        shutil.rmtree(tmpdir, ignore_errors=True)


# ─── STEP 4: FETCH & SUMMARISE COMMENTS ──────────────────────────────────────

def get_top_comments(youtube, video_id: str, max_comments: int = 100) -> list[dict]:
    """Fetch up to max_comments comments ordered by relevance."""
    comments = []
    try:
        next_page = None
        while len(comments) < max_comments:
            resp = youtube.commentThreads().list(
                part="snippet",
                videoId=video_id,
                order="relevance",
                maxResults=min(100, max_comments - len(comments)),
                pageToken=next_page,
            ).execute()
            for item in resp.get("items", []):
                top = item["snippet"]["topLevelComment"]["snippet"]
                comments.append({
                    "text":  top["textDisplay"],
                    "likes": int(top.get("likeCount", 0)),
                })
            next_page = resp.get("nextPageToken")
            if not next_page:
                break
    except Exception:
        pass  # Comments disabled or quota hit
    return comments


def summarise_comments(comments: list[dict]) -> dict:
    """Return theme counts and top-5 liked comments."""
    theme_counts = {theme: 0 for theme in THEME_MAP}
    for c in comments:
        lower = c["text"].lower()
        for theme, keywords in THEME_MAP.items():
            if any(kw in lower for kw in keywords):
                theme_counts[theme] += 1

    # Keep only themes that actually matched
    active_themes = {t: n for t, n in theme_counts.items() if n > 0}

    top5 = sorted(comments, key=lambda c: c["likes"], reverse=True)[:5]

    return {"themes": active_themes, "top5": top5}


# ─── STEP 5: BUILD .DOCX ──────────────────────────────────────────────────────

BLUE       = RGBColor(0x1a, 0x73, 0xe8)   # Google-blue
LIGHT_BLUE = RGBColor(0xdb, 0xe9, 0xf8)   # Row highlight
DARK_GREY  = RGBColor(0x40, 0x40, 0x40)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc: Document):
    """Insert a thin horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


def _heading(doc: Document, text: str, level: int):
    """Add a heading with custom blue colour."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE
    return h


def build_docx(videos: list[dict], output_path: str):
    """Build the full .docx report."""
    doc = Document()

    # ── Margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("LoadedDiceShorts")
    run.font.size  = Pt(36)
    run.font.bold  = True
    run.font.color.rgb = BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run("YouTube Shorts — Full Content & Comment Analysis")
    run2.font.size  = Pt(16)
    run2.font.color.rgb = DARK_GREY

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated = datetime.now().strftime("%d %b %Y")
    run3 = meta_p.add_run(
        f"Total Shorts: {len(videos)}  |  Generated: {generated}"
    )
    run3.font.size   = Pt(12)
    run3.font.italic = True
    run3.font.color.rgb = DARK_GREY

    doc.add_page_break()

    # ════════════════════════════════════════════
    # PAGE 2 — OVERVIEW TABLE
    # ════════════════════════════════════════════
    _heading(doc, "Overview", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["#", "Title", "Views", "Published"]
    for i, (cell, text) in enumerate(zip(hdr_cells, headers)):
        cell.text = text
        _set_cell_bg(cell, "1a73e8")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold      = True
                run.font.size      = Pt(10)

    for idx, v in enumerate(videos, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = v["title"]
        row_cells[2].text = f"{v['views']:,}"
        row_cells[3].text = v["published"]
        if idx % 2 == 0:
            for cell in row_cells:
                _set_cell_bg(cell, "dbe9f8")
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # ONE SECTION PER VIDEO
    # ════════════════════════════════════════════
    for idx, v in enumerate(videos, 1):
        # H1 — title
        h1 = doc.add_heading(f"#{idx}  {v['title']}", level=1)
        for run in h1.runs:
            run.font.color.rgb = BLUE

        # Stats line
        stats_p = doc.add_paragraph()
        stats_run = stats_p.add_run(
            f"👁 {v['views']:,} views  •  👍 {v['likes']:,} likes  •  "
            f"💬 {v['comment_count']:,} comments  •  📅 {v['published']}  •  "
            f"🔗 {v['url']}"
        )
        stats_run.font.size   = Pt(9)
        stats_run.font.italic = True
        stats_run.font.color.rgb = DARK_GREY

        _add_horizontal_rule(doc)

        # Transcript
        h2 = doc.add_heading("📝 Script / Transcript", level=2)
        for run in h2.runs:
            run.font.color.rgb = BLUE

        tx = v["transcript"] or "[No transcript]"
        tp = doc.add_paragraph(tx)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.size = Pt(10)

        _add_horizontal_rule(doc)

        # Comment Summary
        h2b = doc.add_heading("💬 Comment Summary", level=2)
        for run in h2b.runs:
            run.font.color.rgb = BLUE

        summary = v.get("comment_summary", {})
        themes  = summary.get("themes", {})
        top5    = summary.get("top5", [])

        if themes:
            doc.add_paragraph("Comment themes detected:", style="Intense Quote") \
               .runs[0].font.size if doc.paragraphs[-1].runs else None

            theme_p = doc.add_paragraph()
            for theme, count in sorted(themes.items(), key=lambda x: -x[1]):
                run = theme_p.add_run(f"  {theme}: {count} comments\n")
                run.font.size = Pt(10)
        else:
            doc.add_paragraph("No theme matches found.").runs[0].font.size

        if top5:
            doc.add_paragraph("Top liked comments:", style="Intense Quote")
            for c in top5:
                cp = doc.add_paragraph(style="List Bullet")
                run = cp.add_run(f'"{c["text"][:300]}"  —  👍 {c["likes"]:,}')
                run.font.size   = Pt(9)
                run.font.italic = True
        else:
            nocom = doc.add_paragraph("Comments not available for this video.")
            for run in nocom.runs:
                run.font.size   = Pt(9)
                run.font.italic = True

        # Page break (except after last video)
        if idx < len(videos):
            doc.add_page_break()

    doc.save(output_path)
    print(f"\n✅ Saved: {output_path}")


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    youtube = build("youtube", "v3", developerKey=API_KEY)

    # Step 1 — Get Shorts IDs
    print("\n=== Step 1: Getting Shorts IDs ===")
    video_ids = get_all_shorts(youtube, CHANNEL_HANDLE)
    if not video_ids:
        print("❌ No Shorts found. Exiting.")
        return

    # Step 2 — Fetch metadata
    print(f"\n=== Step 2: Fetching metadata for {len(video_ids)} videos ===")
    videos = fetch_video_metadata(youtube, video_ids)

    # Sort by views descending
    videos.sort(key=lambda v: v["views"], reverse=True)

    # Steps 3 & 4 — Transcripts + comments
    print(f"\n=== Steps 3 & 4: Transcripts + Comments ===")
    for i, v in enumerate(videos, 1):
        print(f"[{i}/{len(videos)}] {v['title'][:60]}")

        v["transcript"] = get_transcript(v["id"])
        time.sleep(0.3)

        comments = get_top_comments(youtube, v["id"], MAX_COMMENTS_PER_VIDEO)
        v["comments"] = comments
        v["comment_summary"] = summarise_comments(comments)
        time.sleep(0.3)

    # Step 5 — Build .docx
    print("\n=== Step 5: Building .docx ===")
    script_dir  = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, OUTPUT_FILE)
    build_docx(videos, output_path)
    print(f"Total videos in report: {len(videos)}")


if __name__ == "__main__":
    main()
